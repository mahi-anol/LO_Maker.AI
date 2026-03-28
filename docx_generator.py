"""
DOCX generator — fills template.docx with AI-generated content.

New architecture (matches pipeline.py per-cell generation):
  Every field in lesson_plan dict maps 1-to-1 to exactly one cell or cell block.
  No section parsing. No fuzzy key matching. Each field is plain text ready to write.

Template cell map:
  TABLE 0 — Header
    R0C1 → teacher_name        R0C3 → subject
    R1C1 → grade               R1C3 → duration

  TABLE 1 — Lesson Vision
    R2C0  → learning_outcome   (plain)
    R4C0  → assess_time + "\n\n" + assess_questions   (plain multi-line)
    R6C0  → assess_exemplar    (numbered prefix bold, rest plain)
    R8C0  → vision block: "What:" + vision_what, "Why (Academic):" + vision_why_ac,
                           "Why (Non-academic):" + vision_why_no
    R8C1  → "How:\n" + key_points
    R10C0 → knowledge_text     (plain, sz=21)
    R12C0 → blooms_skills      (one skill per paragraph, plain)

  TABLE 2 — Lesson Method
    R3C0  → launch_teacher     R3C1 → launch_student    R3C2 → launch_materials
    R6C0  → explore_teacher    R6C1 → explore_student   R6C2 → explore_materials
    R9C0  → concept_teacher    R9C1 → concept_student   R9C2 → concept_materials
    R12C0 → guided_teacher     R12C1 → guided_student   R12C2 → guided_materials
    R15C0 → indep_teacher      R15C1 → indep_student    R15C2 → indep_materials
    R18C0 → closing_teacher    R18C1 → closing_student  R18C2 → closing_materials
"""

import os
import re
import shutil
import copy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATE_PATH = os.path.join(os.path.dirname(__file__), "template.docx")
FONT = "Noto Sans Bengali"


# ══════════════════════════════════════════════════════════════════════════════
# XML helpers
# ══════════════════════════════════════════════════════════════════════════════

def _get_ref_pPr(cell):
    """Return a deep copy of the first paragraph's pPr (for consistent spacing)."""
    paras = cell._tc.findall(qn("w:p"))
    if paras:
        pPr = paras[0].find(qn("w:pPr"))
        if pPr is not None:
            return copy.deepcopy(pPr)
    return None


def _clear_cell(cell):
    """Remove all content from cell; keep one empty paragraph shell."""
    tc = cell._tc
    paras = tc.findall(qn("w:p"))
    for p in paras[1:]:
        tc.remove(p)
    if paras:
        p0 = paras[0]
        for child in list(p0):
            if child.tag in (
                qn("w:r"), qn("w:hyperlink"),
                qn("w:bookmarkStart"), qn("w:bookmarkEnd"),
                qn("w:ins"), qn("w:del"),
            ):
                p0.remove(child)


def _make_rPr(bold: bool, font: str, sz_half: int = None) -> OxmlElement:
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:cs"), font)
    rPr.append(rFonts)
    if bold:
        rPr.append(OxmlElement("w:b"))
        rPr.append(OxmlElement("w:bCs"))
    if sz_half:
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(sz_half))
        rPr.append(sz)
        szCs = OxmlElement("w:szCs")
        szCs.set(qn("w:val"), str(sz_half))
        rPr.append(szCs)
    return rPr


def _make_run(text: str, bold: bool, font: str = FONT,
              sz_half: int = None) -> OxmlElement:
    run = OxmlElement("w:r")
    run.append(_make_rPr(bold, font, sz_half))
    t = OxmlElement("w:t")
    t.text = text if text else " "
    if not text or text != text.strip() or not text.strip():
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    run.append(t)
    return run


def _new_para(tc, ref_pPr=None) -> OxmlElement:
    p = OxmlElement("w:p")
    if ref_pPr is not None:
        p.append(copy.deepcopy(ref_pPr))
    tc.append(p)
    return p


# ══════════════════════════════════════════════════════════════════════════════
# Bold detection helpers (for method cells — Teacher/Student Action columns)
# ══════════════════════════════════════════════════════════════════════════════

# Patterns that mark a line as a bold sub-header inside method cells,
# matching the formatting in the example LP_7_Math_01-03-26.docx
_BOLD_HEADER_RE = re.compile(
    r'^('
    r'General Greeting'
    r'|SEL\b.*'
    r'|Sparking Curiosity.*'
    r'|Revealing LO.*'
    r'|Why\s*/\s*Relevance.*'
    r'|Classroom Rules Reminder.*'
    r'|Prior Knowledge Check.*'
    r'|Transition to Activity.*'
    r'|Teacher Feedback.*'
    r'|Teacher Action'
    r'|Pictorial\s*/\s*Representation.*'
    r'|Exit Ticket.*'
    r'|Suggested Time:?\s*\d+.*'
    r'|Problem\s+\d+\s*:.*'
    r'|Task\s+\d+\s*:.*'
    r'|উদাহরণ\s*\d+\s*:.*'
    r'|Teacher Say\s*:?'
    r'|Teacher Do\s*:?'
    r'|Check for Understanding\s*:?'
    r'|WTD\s+নির্দেশনা\s*:?'
    r'|সমাধান\s*:?'
    r'|সমস্যা\s+\d+\s*:?'
    r')$',
    re.IGNORECASE
)


def _is_bold_header(line: str) -> bool:
    """Return True if this line should be rendered as bold in a method cell."""
    return bool(_BOLD_HEADER_RE.match(line.strip()))


def _parse_method_cell(text: str) -> list:
    """
    Parse a Teacher/Student/Materials cell into bold-aware paragraph list.
    Sub-section labels matching _BOLD_HEADER_RE → bold; all else → plain.
    """
    paragraphs = []
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            paragraphs.append([("", False)])
        elif _is_bold_header(stripped):
            paragraphs.append([(stripped, True)])
        else:
            paragraphs.append([(stripped, False)])
    return paragraphs if paragraphs else [[("", False)]]


# ══════════════════════════════════════════════════════════════════════════════
# Cell writers
# ══════════════════════════════════════════════════════════════════════════════

def write_plain(cell, text: str, sz_half: int = None):
    """Write plain (non-bold) multi-line text. Each line → one paragraph."""
    ref_pPr = _get_ref_pPr(cell)
    _clear_cell(cell)
    tc = cell._tc
    lines = text.splitlines() if text and text.strip() else [""]
    for i, line in enumerate(lines):
        if i == 0:
            paras = tc.findall(qn("w:p"))
            p = paras[0] if paras else _new_para(tc, ref_pPr)
        else:
            p = _new_para(tc, ref_pPr)
        p.append(_make_run(line, bold=False, sz_half=sz_half))


def write_paragraphs(cell, paragraphs: list):
    """
    Write structured paragraphs.
    paragraphs = list of lists of (text, bold) run tuples.
    Each inner list = one paragraph.
    """
    ref_pPr = _get_ref_pPr(cell)
    _clear_cell(cell)
    tc = cell._tc
    if not paragraphs:
        paragraphs = [[("", False)]]
    for i, runs in enumerate(paragraphs):
        if i == 0:
            existing = tc.findall(qn("w:p"))
            p = existing[0] if existing else _new_para(tc, ref_pPr)
        else:
            p = _new_para(tc, ref_pPr)
        for text, bold in runs:
            p.append(_make_run(text, bold=bold))


# ══════════════════════════════════════════════════════════════════════════════
# Content builders — each takes exactly the field(s) it needs
# ══════════════════════════════════════════════════════════════════════════════

def _lines_to_plain_paragraphs(text: str) -> list:
    """Convert plain multi-line text into paragraph list for write_paragraphs."""
    lines = [l for l in text.splitlines() if l.strip()]
    if not lines:
        return [[("", False)]]
    return [[(line, False)] for line in lines]


def build_assessment_cell(assess_time: str, assess_questions: str) -> list:
    """
    R4: time line (plain) + blank line + questions (plain).
    """
    paragraphs = []
    time_line = assess_time.strip() if assess_time.strip() else "সময়: ৫ মিনিট | পূর্ণমান: ৬ মার্ক"
    paragraphs.append([(time_line, False)])
    paragraphs.append([("", False)])
    for line in assess_questions.splitlines():
        if line.strip():
            paragraphs.append([(line, False)])
    return paragraphs


def build_exemplar_paragraphs(exemplar: str) -> list:
    """
    R6: numbered prefix bold (e.g. "১) "), rest plain — matching example doc.
    """
    paragraphs = []
    for line in exemplar.splitlines():
        line = line.strip()
        if not line:
            continue
        m = re.match(r'^([১২৩৪৫৬৭৮৯\d]+[)\.\|।]\s*)', line)
        if m:
            prefix = m.group(1)
            rest = line[len(prefix):]
            paragraphs.append([(prefix, True), (rest, False)])
        else:
            paragraphs.append([(line, False)])
    return paragraphs if paragraphs else [[("", False)]]


def build_vision_cell(vision_what: str, vision_why_ac: str, vision_why_no: str) -> list:
    """
    R8 C0: What / Why (Academic) / Why (Non-academic). All plain.
    """
    paragraphs = []

    def add_block(label: str, content: str):
        if not content.strip():
            return
        paragraphs.append([(label, False)])
        paragraphs.append([("", False)])
        for line in content.splitlines():
            if line.strip():
                paragraphs.append([(line, False)])
        paragraphs.append([("", False)])

    add_block("What:", vision_what)
    add_block("Why (Academic):", vision_why_ac)
    add_block("Why (Non-academic):", vision_why_no)

    return paragraphs if paragraphs else [[("", False)]]


def build_how_cell(key_points: str) -> list:
    """R8 C1: 'How:' header plain, steps plain."""
    paragraphs = [
        [("How:", False)],
        [("", False)],
    ]
    for line in key_points.splitlines():
        if line.strip():
            paragraphs.append([(line, False)])
    return paragraphs


def build_blooms_cell(blooms_skills: str) -> list:
    """R12 C0: one skill per paragraph, plain."""
    lines = [l.strip() for l in blooms_skills.splitlines() if l.strip()]
    if not lines:
        return [[("", False)]]
    return [[(line, False)] for line in lines]


def build_method_cell(text: str) -> list:
    """
    Parse teacher/student action cells with smart bold sub-headers.
    Matches formatting in LP_7_Math_01-03-26.docx example:
      - Sub-section labels (General Greeting, SEL, Problem 1:, Task 1:, etc.) → bold
      - All body text → plain
    """
    return _parse_method_cell(text)


# ══════════════════════════════════════════════════════════════════════════════
# Main generator
# ══════════════════════════════════════════════════════════════════════════════

def generate_docx(lesson_plan: dict, output_path: str) -> str:
    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(
            f"template.docx not found at {TEMPLATE_PATH}. "
            "Place the template file next to this script."
        )

    shutil.copy2(TEMPLATE_PATH, output_path)
    doc = Document(output_path)

    t0 = doc.tables[0]
    t1 = doc.tables[1]
    t2 = doc.tables[2]

    # ══════════════════════════════════════════════════════════════════════════
    # TABLE 0 — Header
    # ══════════════════════════════════════════════════════════════════════════
    write_plain(t0.rows[0].cells[1], lesson_plan.get("teacher_name", ""))
    write_plain(t0.rows[0].cells[3], lesson_plan.get("subject", ""))
    write_plain(t0.rows[1].cells[1], lesson_plan.get("grade", ""))
    write_plain(t0.rows[1].cells[3], lesson_plan.get("duration", ""))

    # ══════════════════════════════════════════════════════════════════════════
    # TABLE 1 — Lesson Vision
    # ══════════════════════════════════════════════════════════════════════════

    # R2 — Learning Outcome
    write_plain(t1.rows[2].cells[0], lesson_plan.get("learning_outcome", ""))

    # R4 — Assessment: time + questions combined
    write_paragraphs(
        t1.rows[4].cells[0],
        build_assessment_cell(
            lesson_plan.get("assess_time", ""),
            lesson_plan.get("assess_questions", ""),
        )
    )

    # R6 — Exemplar (bold numbered prefix)
    write_paragraphs(
        t1.rows[6].cells[0],
        build_exemplar_paragraphs(lesson_plan.get("assess_exemplar", ""))
    )

    # R8 C0 — What + Why block (plain)
    write_paragraphs(
        t1.rows[8].cells[0],
        build_vision_cell(
            lesson_plan.get("vision_what", ""),
            lesson_plan.get("vision_why_ac", ""),
            lesson_plan.get("vision_why_no", ""),
        )
    )

    # R8 C1 — How steps (plain)
    write_paragraphs(
        t1.rows[8].cells[1],
        build_how_cell(lesson_plan.get("key_points", ""))
    )

    # R10 C0 — Knowledge from LO (plain, sz=21)
    write_plain(t1.rows[10].cells[0],
                lesson_plan.get("knowledge_text", ""), sz_half=21)

    # R12 C0 — Bloom's skills (plain, one per paragraph)
    write_paragraphs(
        t1.rows[12].cells[0],
        build_blooms_cell(lesson_plan.get("blooms_skills", ""))
    )
    # R12 C1 — PI/BI: leave template as-is

    # ══════════════════════════════════════════════════════════════════════════
    # TABLE 2 — Lesson Method
    # Fixed rows (0,1,2,4,5,7,8,10,11,13,14,16,17) are NOT touched.
    # ══════════════════════════════════════════════════════════════════════════

    def fill_method_row(row_idx: int, teacher: str, student: str, materials: str):
        row = t2.rows[row_idx]
        write_paragraphs(row.cells[0], build_method_cell(teacher))
        write_paragraphs(row.cells[1], build_method_cell(student))
        # Materials column: plain text (short lists, no sub-headers)
        write_plain(row.cells[2], materials)

    fill_method_row(3,
                    lesson_plan.get("launch_teacher", ""),
                    lesson_plan.get("launch_student", ""),
                    lesson_plan.get("launch_materials", ""))

    fill_method_row(6,
                    lesson_plan.get("explore_teacher", ""),
                    lesson_plan.get("explore_student", ""),
                    lesson_plan.get("explore_materials", ""))

    fill_method_row(9,
                    lesson_plan.get("concept_teacher", ""),
                    lesson_plan.get("concept_student", ""),
                    lesson_plan.get("concept_materials", ""))

    fill_method_row(12,
                    lesson_plan.get("guided_teacher", ""),
                    lesson_plan.get("guided_student", ""),
                    lesson_plan.get("guided_materials", ""))

    fill_method_row(15,
                    lesson_plan.get("indep_teacher", ""),
                    lesson_plan.get("indep_student", ""),
                    lesson_plan.get("indep_materials", ""))

    fill_method_row(18,
                    lesson_plan.get("closing_teacher", ""),
                    lesson_plan.get("closing_student", ""),
                    lesson_plan.get("closing_materials", ""))

    doc.save(output_path)
    return output_path